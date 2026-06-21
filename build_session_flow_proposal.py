from docx import Document
from docx.shared import Pt, RGBColor, Inches

doc = Document()

for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

MUTED   = RGBColor(0x88, 0x88, 0x88)
DARK    = RGBColor(0x1A, 0x1A, 0x1A)
MID     = RGBColor(0x44, 0x44, 0x44)
GREEN   = RGBColor(0x1A, 0x7A, 0x3C)
RED     = RGBColor(0x99, 0x22, 0x22)

def title(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(14)
    r.font.color.rgb = DARK

def section_label(text, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(10)
    r.font.color.rgb = color or MUTED

def body(text, italic=False, color=None, space_before=2, space_after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.italic = italic
    if color:
        r.font.color.rgb = color

def divider():
    p = doc.add_paragraph('─' * 62)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(14)
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)


# ── Title ─────────────────────────────────────────────────────
title('Session Script — Proposed Changes')
body('Miracle of Mind · Boom Buddy Resource Kit', italic=True, color=MUTED)
body('For review and approval.', italic=True, color=MUTED)

divider()

# ══════════════════════
# BEFORE
# ══════════════════════
section_label('BEFORE', color=RED)

section_label('WELCOME & INTRODUCTION')
body(
    'Good Morning / Good Evening everyone.\n'
    'My name is ______. I am a student / working professional.\n'
    'I\'m also a volunteer with Isha Foundation founded by Sadhguru. Sadhguru has been offering tools for wellbeing and self-transformation rooted in the science of yoga that have touched the lives of millions of people worldwide over the past four decades.\n'
    'Today, we\'ll be learning a simple yet powerful meditation designed by Sadhguru — that can help you to experience the Miracle of the Mind, to take charge of your mental wellbeing and explore your peak potential.'
)

section_label('SADHGURU & ISHA INTRO  (Optional)')
body(
    'This meditation is designed by Sadhguru, a yogi, and New York Times bestselling author whose work is centered on human wellbeing and inner transformation. He has inspired millions around the world with his indiscriminate involvement with all aspects of life — be it offering tools for transformation or inspiring people around the world to support environmental initiatives or bringing about fundamental changes in education and more. His clarity, wisdom and sense of humor are unmatched. He is an embodiment of a human being at his peak potential and this inspires us to strive for our own growth and transformation.\n\n'
    'Miracle of Mind is one such meditation — designed to be simple, easy to practice, and supportive in everyday situations.\n\n'
    'We would like to share a short video introducing Sadhguru to you.  ▶ Play video'
)

section_label('INTERACTION')
body(
    'Do you feel your thoughts are noisy sometimes?\n'
    'Have you wished you could just mute them for a while so you could focus on something?\n'
    'How many times have we tried to control our thoughts, and it didn\'t work?\n\n'
    '(Pause — allow them to respond)'
)

divider()

# ══════════════════════
# AFTER
# ══════════════════════
section_label('AFTER', color=GREEN)

section_label('WELCOME & INTRODUCTION')
body(
    'Good Morning / Good Evening everyone.\n'
    'My name is ______. I am a student / working professional.\n'
    'I\'m also a volunteer with Isha Foundation founded by Sadhguru. Sadhguru has been offering tools for wellbeing and self-transformation rooted in the science of yoga that have touched the lives of millions of people worldwide over the past four decades.',
    color=DARK
)

section_label('SADHGURU & ISHA INTRO')
body(
    'He is a yogi, New York Times bestselling author, and humanitarian — whose work spans inner transformation, environmental initiatives, and education. His clarity, wisdom, and sense of humor are unmatched. He is an embodiment of a human being at his peak potential, and this inspires us to strive for our own growth and transformation.\n\n'
    'We\'d like to share a short video.  ▶ Play video',
    color=DARK
)

section_label('INTERACTION')
body(
    'Today, we\'ll be learning a simple yet powerful meditation designed by Sadhguru — that can help you to experience the Miracle of the Mind, to take charge of your mental wellbeing and explore your peak potential.\n\n'
    'Do you feel your thoughts are noisy sometimes?\n'
    'Have you wished you could just mute them for a while so you could focus on something?\n'
    'How many times have we tried to control our thoughts, and it didn\'t work?\n\n'
    '(Pause — allow them to respond)',
    color=DARK
)

divider()
body('Please confirm if these changes are approved for implementation.', italic=True, color=MUTED)

# ── Save ──────────────────────────────────────────────────────
out = '/Users/aravindmurari/Downloads/MoM_SessionScript_FlowProposal.docx'
doc.save(out)
print(f'Saved: {out}')
