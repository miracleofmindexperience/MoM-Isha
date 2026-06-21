from docx import Document
from docx.shared import Pt, RGBColor, Inches

doc = Document()

for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

DARK  = RGBColor(0x1A, 0x1A, 0x1A)
MID   = RGBColor(0x44, 0x44, 0x44)
MUTED = RGBColor(0x88, 0x88, 0x88)

def title(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = DARK

def subtitle(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(14)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = MUTED

def section_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = DARK

def body(text, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.italic = italic
    r.font.color.rgb = MID

def bullet(text, italic_part=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0.3)
    if italic_part:
        # Split on italic_part and render it italic
        idx = text.find(italic_part)
        if idx >= 0:
            r1 = p.add_run(text[:idx])
            r1.font.size = Pt(11); r1.font.color.rgb = MID
            r2 = p.add_run(italic_part)
            r2.font.size = Pt(11); r2.italic = True; r2.font.color.rgb = MID
            r3 = p.add_run(text[idx+len(italic_part):])
            r3.font.size = Pt(11); r3.font.color.rgb = MID
            return p
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.color.rgb = MID

def divider():
    p = doc.add_paragraph('─' * 62)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(10)
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)


# ── Title block ───────────────────────────────────────────────
title('Session Media — Behavior Reference')
subtitle('Boom Buddy Resource Kit · Miracle of Mind')
divider()

# ── 1. Default state ──────────────────────────────────────────
section_heading('Default state (no video selected)')
body(
    'When the Session Media card is opened, the video area shows a black screen with the text: '
    '"Select a video from the playlist to play." No video loads or plays automatically.'
)

divider()

# ── 2. Playing a video ────────────────────────────────────────
section_heading('Playing a video (main view only)')
body(
    'When the volunteer clicks any item in the playlist, the video loads and begins playing '
    'directly in the main page. The playlist item is highlighted to show what is currently '
    'selected. The "Now Playing" label below the video area updates with the video title and '
    'context (e.g., "Pre-session · ~13 min").'
)

divider()

# ── 3. Opening Presentation View ──────────────────────────────
section_heading('Opening Presentation View')
body(
    'The "Open Presentation View" button opens present.html in a separate browser window — '
    'intended to be sent to a projector or second monitor.'
)
body('Two scenarios:', italic=False)

p = doc.add_paragraph(style='List Number')
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after  = Pt(4)
p.paragraph_format.left_indent  = Inches(0.3)
r = p.add_run('No video selected yet')
r.bold = True; r.font.size = Pt(11); r.font.color.rgb = MID
r2 = p.add_run(
    ' — The presentation window opens and waits. When the volunteer clicks a playlist item, '
    'the video plays only on the presentation screen. The main page shows an overlay: '
    '"Video is playing on the presentation screen. Close the presentation view to preview here."'
)
r2.font.size = Pt(11); r2.font.color.rgb = MID

p2 = doc.add_paragraph(style='List Number')
p2.paragraph_format.space_before = Pt(4)
p2.paragraph_format.space_after  = Pt(8)
p2.paragraph_format.left_indent  = Inches(0.3)
r3 = p2.add_run('A video is already playing on the main page')
r3.bold = True; r3.font.size = Pt(11); r3.font.color.rgb = MID
r4 = p2.add_run(
    ' — The moment the presentation window opens, local playback is immediately stopped and '
    'the same overlay appears on the main page. The currently selected video is automatically '
    'pushed to and begins playing on the presentation screen.'
)
r4.font.size = Pt(11); r4.font.color.rgb = MID

body(
    'In both cases, while the presentation window is open, clicking any playlist item on the '
    'main page routes the video to the presentation screen only — the main page continues to '
    'show the overlay.'
)

divider()

# ── 4. Closing Presentation View ─────────────────────────────
section_heading('Closing Presentation View')
body(
    'When the presentation window is closed, the main page automatically detects it (within '
    'half a second) and returns to the idle state — the overlay clears and the '
    '"Select a video from the playlist to play" placeholder reappears. No video plays '
    'automatically; the volunteer can then click a playlist item to preview locally.'
)

divider()

# ── 5. Offline ────────────────────────────────────────────────
section_heading('Offline / Download for offline use')
body(
    'There is a separate "Download for offline use" panel. Volunteers can pre-download '
    'individual videos to their browser cache (stored up to 7 days). When a cached video '
    'is available, it plays from the local cache instead of streaming from Dropbox — '
    'useful for sessions with unreliable internet.'
)

# ── Save ──────────────────────────────────────────────────────
out = '/Users/aravindmurari/Downloads/MoM_SessionMedia_Behavior.docx'
doc.save(out)
print(f'Saved: {out}')
